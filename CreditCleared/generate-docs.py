#!/usr/bin/env python3
"""Credit Cleared document generator (spec Section 6).

Reads job.json + consumer-analysis-response.json from a job folder and
produces five branded .docx files in <job-folder>/docs/:

    1. Credit Cleared Roadmap.docx
    2. Call Scripts.docx
    3. Dispute Letters.docx
    4. Goodwill Letters.docx
    5. Debt Validation Letters.docx

The spec's Document 1 mapping (Account Summary, Priority Action Plan,
Creditor Contact Directory, Milestone Map, Educational Summary) predates
CC-Consumer-Analysis-Prompt.docx's 11-section version and doesn't mention
where the two newer sections (Inquiry Cross-Reference Analysis, Personal
Credit Profile Building) or the Executive Summary go. This script places
all of the analytical/educational sections in Document 1 alongside the
ones the spec named, and keeps Documents 2-5 as the individually-mailable
letter/script sets -- see SECTION_TO_DOCUMENT below.

Formatting note: the spec asks for "highlighted key phrases" and
"response handling in italics" in Call Scripts.docx. The Claude response
is unstructured prose, not per-field JSON, so there's no reliable marker
to key that off. What's implemented instead: short, colon-terminated or
title-case lines are rendered as bold sub-headings; everything else is a
plain paragraph. That's a deliberate simplification of the spec's literal
formatting request, not a full implementation of it.

Document 6 (Funding Roadmap.docx) is NOT generated here -- per the
upsell-only decision, it's produced by a separate on-demand script when a
client purchases the $197 Business Funding add-on.
"""

import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x0A, 0x16, 0x28)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
LEGAL_FOOTER = "Credit Cleared provides credit report analysis, dispute preparation, and creditor negotiation guidance. Individual results vary and are not guaranteed."

SECTION_HEADERS = [
    "EXECUTIVE SUMMARY",
    "SECTION 1: ACCOUNT SUMMARY",
    "SECTION 2: PRIORITY ACTION PLAN",
    "SECTION 3: CREDITOR CONTACT DIRECTORY",
    "SECTION 4: WORD-FOR-WORD CALL SCRIPTS",
    "SECTION 5: DISPUTE LETTERS",
    "SECTION 6: GOODWILL LETTERS",
    "SECTION 7: DEBT VALIDATION LETTERS",
    "SECTION 8: INQUIRY CROSS-REFERENCE ANALYSIS",
    "SECTION 9: MONTH-BY-MONTH MILESTONE MAP",
    "SECTION 10: PERSONAL CREDIT PROFILE BUILDING",
]

DOCUMENTS = [
    {
        "filename": "Credit Cleared Roadmap.docx",
        "sections": [
            "EXECUTIVE SUMMARY",
            "SECTION 1: ACCOUNT SUMMARY",
            "SECTION 2: PRIORITY ACTION PLAN",
            "SECTION 3: CREDITOR CONTACT DIRECTORY",
            "SECTION 8: INQUIRY CROSS-REFERENCE ANALYSIS",
            "SECTION 9: MONTH-BY-MONTH MILESTONE MAP",
            "SECTION 10: PERSONAL CREDIT PROFILE BUILDING",
        ],
    },
    {
        "filename": "Call Scripts.docx",
        "sections": ["SECTION 4: WORD-FOR-WORD CALL SCRIPTS"],
    },
    {
        "filename": "Dispute Letters.docx",
        "sections": ["SECTION 5: DISPUTE LETTERS"],
    },
    {
        "filename": "Goodwill Letters.docx",
        "sections": ["SECTION 6: GOODWILL LETTERS"],
    },
    {
        "filename": "Debt Validation Letters.docx",
        "sections": ["SECTION 7: DEBT VALIDATION LETTERS"],
    },
]


def parse_sections(raw_text: str) -> dict:
    """Split the Claude response into {header: body} using the known,
    exact section headers the prompt instructed it to use."""
    pattern = "|".join(re.escape(h) for h in SECTION_HEADERS)
    regex = re.compile(rf"^\s*#*\s*\**\s*({pattern})\s*\**\s*#*\s*$", re.MULTILINE)

    matches = list(regex.finditer(raw_text))
    sections = {}
    for i, m in enumerate(matches):
        header = m.group(1)
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(raw_text)
        sections[header] = raw_text[start:end].strip()
    return sections


def is_subheading(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if len(stripped) > 70:
        return False
    if stripped.endswith(":"):
        return True
    if stripped.isupper() and len(stripped.split()) <= 8:
        return True
    return False


def set_cell_shading(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_branded_header(doc: Document, title: str, client_name: str, generated_date: str):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "0A1628")
    p = cell.paragraphs[0]
    run = p.add_run("CREDIT CLEARED")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p2 = cell.add_paragraph()
    run2 = p2.add_run(title)
    run2.bold = True
    run2.font.size = Pt(13)
    run2.font.color.rgb = GOLD

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Prepared for {client_name}  |  {generated_date}")
    meta_run.italic = True
    meta_run.font.size = Pt(10)
    doc.add_paragraph()


def add_body_text(doc: Document, text: str):
    for raw_line in text.split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue
        # Strip markdown artifacts -- the doc has its own bold/heading
        # styling, so literal ## and ** characters shouldn't show up.
        cleaned = re.sub(r"^#+\s*", "", line.strip())
        cleaned = cleaned.replace("**", "")
        p = doc.add_paragraph()
        run = p.add_run(cleaned)
        if is_subheading(cleaned):
            run.bold = True
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(11)


def fix_zoom_element(doc: Document):
    """python-docx's bundled default.docx template writes <w:zoom w:val="..."/>
    without the w:percent attribute the OOXML schema requires. Word tolerates
    it, but stricter readers don't -- cheap to fix before saving."""
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")


def add_legal_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(LEGAL_FOOTER)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def build_document(out_path: Path, title: str, client_name: str, generated_date: str, body_sections: list):
    doc = Document()
    add_branded_header(doc, title, client_name, generated_date)
    for header, body in body_sections:
        if not body:
            continue
        heading = doc.add_paragraph()
        heading_run = heading.add_run(header)
        heading_run.bold = True
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = NAVY
        add_body_text(doc, body)
        doc.add_paragraph()
    add_legal_footer(doc)
    fix_zoom_element(doc)
    doc.save(str(out_path))


def main():
    parser = argparse.ArgumentParser(description="Generate Credit Cleared roadmap documents from a Claude analysis response.")
    parser.add_argument("--job-folder", required=True, help="Path to the job folder (contains job.json, consumer-analysis-response.json)")
    parser.add_argument("--job-id", required=True, help="Job ID, for logging")
    args = parser.parse_args()

    job_folder = Path(args.job_folder)
    job_json_path = job_folder / "job.json"
    response_path = job_folder / "consumer-analysis-response.json"

    if not job_json_path.exists():
        print(f"ERROR: job.json not found at {job_json_path}", file=sys.stderr)
        return 1
    if not response_path.exists():
        print(f"ERROR: consumer-analysis-response.json not found at {response_path}", file=sys.stderr)
        return 1

    job = json.loads(job_json_path.read_text(encoding="utf-8"))
    response = json.loads(response_path.read_text(encoding="utf-8"))

    client = job.get("client", {})
    client_name = f"{client.get('first_name', '')} {client.get('last_name', '')}".strip() or "Client"
    generated_date = datetime.now().strftime("%B %d, %Y")

    raw_text = response.get("raw_text", "")
    if not raw_text.strip():
        print("ERROR: consumer-analysis-response.json has empty raw_text", file=sys.stderr)
        return 1

    sections = parse_sections(raw_text)
    missing = [h for h in SECTION_HEADERS if h not in sections]
    if missing:
        print(f"WARNING: could not locate {len(missing)} expected section(s) in the response: {', '.join(missing)}")

    docs_dir = job_folder / "docs"
    docs_dir.mkdir(exist_ok=True)

    for doc_spec in DOCUMENTS:
        body_sections = [(h, sections.get(h, "")) for h in doc_spec["sections"]]
        out_path = docs_dir / doc_spec["filename"]
        doc_title = doc_spec["filename"].replace(".docx", "")
        build_document(out_path, doc_title, client_name, generated_date, body_sections)
        print(f"Wrote {out_path}")

    print(f"Generated {len(DOCUMENTS)} documents for job {args.job_id} in {docs_dir}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
